"""
Generates the IRB consent form as a .docx file for submission.
Run from the project root: python scripts/generate_consent_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "IRB_Consent_Form_ScholarScaffold.docx"
)


def set_font(run, name="Times New Roman", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12)
    return p


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def build_doc():
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ── HEADER ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("TOWSON UNIVERSITY")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Institutional Review Board")
    set_font(run, size=12, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("INFORMED CONSENT / STUDY INFORMATION SHEET")
    set_font(run, size=12, bold=True)

    add_horizontal_rule(doc)

    # Study info block
    info_lines = [
        ("Study Title:", "Evaluating the Effectiveness of ScholarScaffold: "
         "AI-Guided Research Literacy and Scholarly Proposal Development Platform"),
        ("Review Type:", "Expedited Review — Category 7"),
        ("Principal Investigator:", "Wendy Whitner, PhD, MPH, LSSBB"),
        ("Co-Investigator:", "Mona A. Mohamed, PhD"),
        ("Institution:", "Towson University, Department of Health Professions"),
        ("IRB Protocol #:", "_______________________ (assigned upon approval)"),
        ("Version Date:", "May 10, 2026"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}  ")
        set_font(r1, size=12, bold=True)
        r2 = p.add_run(value)
        set_font(r2, size=12)

    add_horizontal_rule(doc)

    # ── SECTION 1 ────────────────────────────────────────────────────────────
    add_heading(doc, "1.  Purpose of This Study")
    add_body(doc,
        "This document provides information about a research study being conducted by faculty at Towson "
        "University. Please read it carefully before deciding whether to participate. You are being invited "
        "to participate because you are enrolled in a course that uses ScholarScaffold, a web-based platform "
        "designed to guide students through the scholarly proposal development process.")
    add_body(doc,
        "The purpose of this research study is to evaluate whether structured, AI-guided scaffolding through "
        "ScholarScaffold produces measurable improvement in the quality of scholarly research proposals, as "
        "reflected in rubric scores from a first draft to a revised draft. The study also examines how "
        "students engage with each stage of the platform and what they think about the experience.")
    add_body(doc,
        "This study involves research. Your participation is completely voluntary. Your decision to participate "
        "or not participate will have no effect on your grade, your standing in this course, or your relationship "
        "with your instructor or Towson University.")

    # ── SECTION 2 ────────────────────────────────────────────────────────────
    add_heading(doc, "2.  What You Will Be Asked to Do")
    add_body(doc,
        "All students in this course will use ScholarScaffold as a course resource to complete their scholarly "
        "proposal assignment. The platform guides you through six sequential stages, each of which must be "
        "completed before moving to the next:")
    for stage in [
        "Stage 1: Research Strategy Coach",
        "Stage 2: Research Design Literacy Module",
        "Stage 3: Article Review Builder",
        "Stage 4: Annotated Bibliography Builder",
        "Stage 5: Proposal Builder",
        "Stage 6: Rubric Scoring Engine",
    ]:
        add_bullet(doc, stage)
    add_body(doc,
        "If you agree to participate in this research study, your platform activity will be included in the "
        "research dataset. This includes your rubric scores on Draft 1 and Draft 2, the time you spent in "
        "each stage, which stages you completed, and your responses to a brief post-experience survey "
        "described below.")
    add_body(doc,
        "At the end of your platform experience, you will be asked to complete a short survey of approximately "
        "5 to 8 questions. The survey asks for your ratings of how easy the platform was to use and how "
        "valuable you found the learning experience, along with one or two open-ended questions asking you to "
        "describe your experience in your own words. The survey should take no more than 10 minutes to complete. "
        "You may skip any survey questions you prefer not to answer.")
    add_body(doc,
        "If you choose not to participate in the research, you will still use ScholarScaffold for your "
        "coursework in exactly the same way. The only difference is that your data will not be included in "
        "the research dataset.")

    # ── SECTION 3 ────────────────────────────────────────────────────────────
    add_heading(doc, "3.  Time Commitment")
    add_body(doc,
        "Your total time commitment for research-specific activities — meaning the consent process and the "
        "post-experience survey — is estimated at 10 to 15 minutes. Your engagement with ScholarScaffold "
        "itself is part of your normal coursework and is not additional time required for research participation.")

    # ── SECTION 4 ────────────────────────────────────────────────────────────
    add_heading(doc, "4.  Who Can Participate")
    add_body(doc, "To participate in this study, you must meet all of the following criteria:")
    for criterion in [
        "Be 18 years of age or older",
        "Be currently enrolled in this course",
        "Be able to provide your own informed consent",
    ]:
        add_bullet(doc, criterion)
    add_body(doc,
        "There are no other eligibility requirements. All enrolled students have equal access to ScholarScaffold "
        "as a course tool regardless of whether they choose to participate in the research.")

    # ── SECTION 5 ────────────────────────────────────────────────────────────
    add_heading(doc, "5.  Risks and Discomforts")
    add_body(doc,
        "The risks associated with this study are minimal. There are no physical risks, no medical procedures, "
        "and no collection of sensitive personal information such as health records or financial information. "
        "The primary risk is loss of confidentiality, meaning that your identity could potentially be linked "
        "to your research data. This risk is minimized through the confidentiality protections described in "
        "Section 7 below.")
    add_body(doc,
        "You may also experience mild discomfort if you feel concerned about your platform performance being "
        "used for research purposes. To address this, please know that the research investigator will not "
        "review or access any participant data until after your final course grade has been submitted. Your "
        "platform work cannot influence your grade in any way.")
    add_body(doc,
        "If you become upset or distressed for any reason, whether related to this study or not, you are "
        "encouraged to contact the Towson University Counseling Center at towson.edu/counseling. If you are "
        "in crisis outside of regular business hours, the 988 Suicide and Crisis Lifeline is available "
        "24 hours a day, 7 days a week by calling or texting 988.")

    # ── SECTION 6 ────────────────────────────────────────────────────────────
    add_heading(doc, "6.  Anticipated Benefits")
    add_body(doc,
        "You are not expected to benefit directly from your participation in this research study. However, by "
        "using ScholarScaffold as a course tool, you may experience improvement in your proposal writing skills "
        "and research literacy, regardless of whether you consent to the research.")
    add_body(doc,
        "The findings of this study will benefit future students and health professions education programs by "
        "contributing to knowledge about whether AI-guided scaffolding can effectively support the development "
        "of scholarly research skills.")

    # ── SECTION 7 ────────────────────────────────────────────────────────────
    add_heading(doc, "7.  Confidentiality")
    add_body(doc,
        "Your participation in this study is confidential, not anonymous. This means that the research team "
        "will be able to link your name to your data during the study period, but every reasonable step will "
        "be taken to protect your privacy. Specifically, the following protections are in place:")
    for bullet in [
        "Upon consenting, you will be assigned a unique coded identifier. All research analyses will use this "
        "coded identifier, not your name.",
        "The file linking your name to your coded identifier will be stored separately from the research dataset "
        "in a password-protected location accessible only to the Principal Investigator and Co-Investigator.",
        "This linkage file will be permanently destroyed within two years of the completion of data collection.",
        "Research data will be retained for a minimum of three years following publication or presentation of "
        "results, in accordance with Towson University research data retention policy, and then securely destroyed.",
        "No identifying information will be included in any publication, conference presentation, or report "
        "resulting from this study.",
        "The Principal Investigator, who is also your course instructor, will not access or review any participant "
        "data until after your final grade for this course has been submitted.",
        "Your survey responses, including any open-ended written responses, will be reviewed and analyzed only "
        "in de-identified form.",
    ]:
        add_bullet(doc, bullet)
    add_body(doc,
        "Please be aware that although every effort will be made to protect your confidentiality, absolute "
        "confidentiality cannot be guaranteed.")

    # ── SECTION 8 ────────────────────────────────────────────────────────────
    add_heading(doc, "8.  Voluntary Participation and Your Rights")
    add_body(doc,
        "Your participation in this research study is completely voluntary. It is your decision whether to "
        "participate. You may choose not to participate, and you may withdraw your consent and discontinue "
        "participation at any time, without penalty, without loss of any benefits to which you are already "
        "entitled, and without negatively affecting your relationship with Towson University, your course "
        "instructor, or your standing in this course.")
    add_body(doc,
        "You do not have to answer any questions you do not want to answer, including any survey items, for "
        "any reason. If you choose to withdraw from the study, data collected up to the point of your "
        "withdrawal will be excluded from the research dataset upon your written or emailed request to the "
        "Principal Investigator.")
    add_body(doc,
        "The Principal Investigator reserves the right to discontinue your participation in the study at any "
        "time if it becomes necessary to protect your welfare or the integrity of the research.")

    # ── SECTION 9 ────────────────────────────────────────────────────────────
    add_heading(doc, "9.  Compensation")
    add_body(doc,
        "You will receive no compensation for your participation in this study. No course credit, extra credit, "
        "or other academic or financial incentive is offered for research participation.")

    # ── SECTION 10 ────────────────────────────────────────────────────────────
    add_heading(doc, "10.  Data Security")
    add_body(doc,
        "All research data are stored on a secured server with role-based access controls, in accordance with "
        "Towson University Office of Technology Services Data Use Standards. Data will not be transmitted "
        "through unsecured channels or stored on personal devices. The platform uses encrypted communications "
        "(HTTPS) for all data transfers.")

    # ── SECTION 11 ────────────────────────────────────────────────────────────
    add_heading(doc, "11.  Contact Information")
    add_body(doc,
        "If you have questions about this study, wish to request a copy of your data, or wish to withdraw your "
        "consent, please contact the research team. If you have questions about your rights as a research "
        "participant or concerns about the conduct of this study, please contact the Towson University IRB "
        "Chair.")

    # Contact block — PI
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Principal Investigator")
    set_font(r, size=12, bold=True)
    for line in [
        "Wendy Whitner, PhD, MPH, LSSBB",
        "Department of Health Professions, Towson University",
        "Phone: (410) 271-3966",
        "Email: wwhitner@towson.edu",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        set_font(r, size=12)

    doc.add_paragraph()

    # Contact block — Co-I
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Co-Investigator")
    set_font(r, size=12, bold=True)
    for line in [
        "Mona A. Mohamed, PhD",
        "Department of Health Professions, Towson University",
        "Phone: (410) 704-4376",
        "Email: mmohamed@towson.edu",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        set_font(r, size=12)

    doc.add_paragraph()

    # Contact block — IRB
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("IRB Chair (for questions about your rights as a participant)")
    set_font(r, size=12, bold=True)
    for line in [
        "Dr. Elizabeth Katz",
        "Office of Sponsored Programs and Research, Towson University",
        "Phone: (410) 704-2236",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        set_font(r, size=12)

    add_horizontal_rule(doc)

    # ── SECTION 12 — ONLINE CONSENT STATEMENT ────────────────────────────────
    add_heading(doc, "12.  Online Consent Statement")
    add_body(doc,
        "Because this consent is obtained online, clicking YES within the ScholarScaffold platform serves as "
        "your electronic signature and constitutes your informed consent to participate. Electronic consent "
        "has the same legal effect as a handwritten signature.")
    add_body(doc, "By clicking YES, you are confirming that:")
    for item in [
        "You are being asked to participate in a research study.",
        "Your participation is completely voluntary and you may withdraw at any time without penalty.",
        "You do not have to answer any questions you do not want to answer.",
        "Your decision to participate or not participate will have no effect on your course grade or your "
        "standing with Towson University.",
        "You have had the opportunity to read this information sheet and ask questions before deciding.",
        "You are 18 years of age or older.",
        "You are able to provide your own informed consent.",
    ]:
        add_bullet(doc, item)

    add_body(doc,
        "A copy of this consent form is available within the ScholarScaffold platform at any time via the "
        "Print or Save option. You may also request a copy from the Principal Investigator.")

    add_horizontal_rule(doc)

    # ── PAPER SIGNATURE BLOCK (for IRB submission) ───────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Participant Signature Block")
    set_font(r, size=12, bold=True)

    add_body(doc,
        "Note: The signature block below is included for IRB submission purposes and for any participants who "
        "require a paper copy. In standard study administration, consent is obtained electronically via the "
        "ScholarScaffold platform (Section 12 above).")

    doc.add_paragraph()

    sig_lines = [
        "Participant Name (print):  ___________________________________________",
        "",
        "Participant Signature:  _____________________________________________  Date:  __________",
        "",
        "Investigator/Research Staff Signature:  ______________________________  Date:  __________",
    ]
    for line in sig_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(line)
        set_font(r, size=12)

    # ── FOOTER NOTE ──────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "IRB Consent Form — ScholarScaffold Study  |  Version Date: May 10, 2026  |  "
        "Towson University IRB Expedited Review — Category 7"
    )
    set_font(r, size=10)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
